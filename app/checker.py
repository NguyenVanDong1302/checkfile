import re
import io
import unicodedata
from typing import Dict, List, Tuple, Any, Iterable

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn


# =========================
# Regex + helpers
# =========================
WORD_RE = re.compile(r"[A-Za-zÀ-ỹĐđ]+", re.UNICODE)
SENT_CAP_RE = re.compile(r"([.!?…])(\s+)([a-zà-ỹđ])", re.UNICODE)

FORBIDDEN_WORDS_VI = ["qá", "cuyện", "táon"]

RULE_BASE = re.compile(r"dự\s*toán\s*mua\s*sắm", re.IGNORECASE | re.UNICODE)
RULE_EXPECT_AFTER = re.compile(r"^\s*:\s*mua\s*sắm\b", re.IGNORECASE | re.UNICODE)


def io_bytes(b: bytes) -> io.BytesIO:
    return io.BytesIO(b)


def _safe_text(s: str) -> str:
    return (s or "").replace("\u00a0", " ").strip()


def _has_page_break_in_paragraph(p: Paragraph) -> bool:
    """
    Detect explicit page breaks: <w:br w:type="page"/>
    """
    try:
        brs = p._p.xpath(".//w:br[@w:type='page']")
        return len(brs) > 0
    except Exception:
        return False


def iter_block_items(doc: Document) -> Iterable[Tuple[str, Any]]:
    """
    Yield ("paragraph", Paragraph) or ("table", Table) in document order.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("paragraph", Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            yield ("table", Table(child, doc))


# =========================
# Heading + capitalization rules
# =========================
def _is_heading_like(text: str, style_name: str | None = None) -> bool:
    """
    Heading if:
    - style name contains "Heading"
    - OR line is almost/all uppercase letters
    """
    if style_name and "heading" in style_name.lower():
        return True

    letters = [ch for ch in (text or "") if ch.isalpha()]
    if len(letters) < 6:
        return False
    upper = sum(1 for ch in letters if ch.isupper())
    lower = sum(1 for ch in letters if ch.islower())
    return lower == 0 and upper > 0


def _is_title_case_word(word: str) -> bool:
    letters = [ch for ch in word if ch.isalpha()]
    if not letters:
        return False
    if not letters[0].isupper():
        return False
    return all(ch.islower() for ch in letters[1:])


def find_weird_mixed_case_words(text: str) -> List[Dict[str, Any]]:
    """
    Flag words like: SƠn / sƠn / SoN ... but allow normal Title Case: Sơn
    """
    hits = []
    s = text or ""
    for m in WORD_RE.finditer(s):
        token = m.group(0)
        letters = [ch for ch in token if ch.isalpha()]
        if len(letters) < 2:
            continue

        has_upper = any(ch.isupper() for ch in letters)
        has_lower = any(ch.islower() for ch in letters)

        if has_upper and has_lower:
            if not _is_title_case_word(token):
                start, end = m.start(), m.end()
                line_no = s.count("\n", 0, start) + 1
                snippet = s[max(0, start - 40):min(len(s), end + 40)].replace("\n", " ↵ ")
                hits.append({
                    "word": token,
                    "line": line_no,
                    "start": start,
                    "end": end,
                    "snippet": snippet,
                    "message": "Từ có kiểu HOA/thường bất thường (ví dụ SƠn).",
                })
    return hits


def find_sentence_start_not_capitalized(text: str) -> List[Dict[str, Any]]:
    """
    After . ! ? … + spaces, the next letter should be uppercase.
    If next is lowercase => flag.
    """
    hits = []
    s = text or ""
    for m in SENT_CAP_RE.finditer(s):
        start = m.start(3)
        end = m.end(3)
        line_no = s.count("\n", 0, start) + 1
        snippet = s[max(0, m.start() - 40):min(len(s), m.end() + 40)].replace("\n", " ↵ ")
        hits.append({
            "char": s[start:end],
            "line": line_no,
            "start": start,
            "end": end,
            "snippet": snippet,
            "message": "Sau dấu chấm/câu hỏi/cảm thán nên viết HOA chữ cái đầu câu.",
        })
    return hits


# =========================
# Phrase scan
# =========================
def _scan_text_for_phrases(
    text: str,
    phrases: List[str],
    *,
    case_sensitive: bool,
    whole_word: bool,
) -> Dict[str, List[Dict[str, Any]]]:
    results: Dict[str, List[Dict[str, Any]]] = {p: [] for p in phrases}
    if not text:
        return results

    flags = 0 if case_sensitive else re.IGNORECASE

    for phrase in phrases:
        if not phrase:
            continue

        escaped = re.escape(phrase)
        if whole_word:
            pattern = r"\b" + escaped + r"\b"
        else:
            pattern = escaped

        for m in re.finditer(pattern, text, flags=flags):
            start, end = m.start(), m.end()
            line_no = text.count("\n", 0, start) + 1

            left = max(0, start - 40)
            right = min(len(text), end + 40)
            snippet = text[left:right].replace("\n", " ↵ ")

            results[phrase].append({"start": start, "end": end, "line": line_no, "snippet": snippet})

    return results


# =========================
# Vietnamese fuzzy spellcheck
# =========================
def vi_strip_tone(s: str) -> str:
    s = s.replace("đ", "d").replace("Đ", "D")
    nfkd = unicodedata.normalize("NFD", s)
    return "".join(ch for ch in nfkd if unicodedata.category(ch) != "Mn")


def damerau_levenshtein_limited(a: str, b: str, max_dist: int) -> int:
    if a == b:
        return 0
    la, lb = len(a), len(b)
    if abs(la - lb) > max_dist:
        return max_dist + 1
    if la == 0:
        return lb
    if lb == 0:
        return la

    prev_prev = list(range(lb + 1))
    prev = [0] * (lb + 1)
    curr = [0] * (lb + 1)

    for i in range(1, la + 1):
        prev[0] = i
        min_in_row = prev[0]
        ai = a[i - 1]

        for j in range(1, lb + 1):
            cost = 0 if ai == b[j - 1] else 1
            deletion = prev_prev[j] + 1
            insertion = prev[j - 1] + 1
            substitution = prev_prev[j - 1] + cost
            v = min(deletion, insertion, substitution)

            if i > 1 and j > 1 and a[i - 1] == b[j - 2] and a[i - 2] == b[j - 1]:
                v = min(v, prev_prev[j - 2] + 1)

            curr[j] = v
            if v < min_in_row:
                min_in_row = v

        if min_in_row > max_dist:
            return max_dist + 1

        prev_prev, prev, curr = prev, curr, prev_prev

    return prev[j]


def find_vi_typos_in_block(expected_word: str, text: str, *, max_dist: int) -> List[Dict[str, Any]]:
    exp = expected_word.strip()
    if not exp or " " in exp:
        return []

    exp_l = exp.lower()
    exp_base = vi_strip_tone(exp_l)

    hits = []
    for m in WORD_RE.finditer(text or ""):
        token = m.group(0)
        tok_l = token.lower()
        if tok_l == exp_l:
            continue

        tok_base = vi_strip_tone(tok_l)
        base_dist = damerau_levenshtein_limited(exp_base, tok_base, max_dist)
        if base_dist > max_dist:
            continue

        dist = damerau_levenshtein_limited(exp_l, tok_l, max_dist)
        if dist <= max_dist:
            start, end = m.start(), m.end()
            line_no = (text or "").count("\n", 0, start) + 1
            left = max(0, start - 40)
            right = min(len(text), end + 40)
            snippet = (text or "")[left:right].replace("\n", " ↵ ")

            hits.append({
                "typed_in_doc": token,
                "distance": dist,
                "line": line_no,
                "start": start,
                "end": end,
                "snippet": snippet,
            })

    hits.sort(key=lambda x: (x["distance"], x["typed_in_doc"].lower()))
    return hits[:50]


# =========================
# Forbidden words + rule "dự toán mua sắm: mua sắm"
# =========================
def _scan_forbidden_words(text: str, forbidden: List[str]) -> List[Dict[str, Any]]:
    if not text:
        return []
    hits = []
    for w in forbidden:
        pattern = re.compile(rf"\b{re.escape(w)}\b", re.IGNORECASE | re.UNICODE)
        for m in pattern.finditer(text):
            start, end = m.start(), m.end()
            line_no = text.count("\n", 0, start) + 1
            left = max(0, start - 40)
            right = min(len(text), end + 40)
            snippet = text[left:right].replace("\n", " ↵ ")
            hits.append({
                "wrong_word": text[start:end],
                "canonical": w,
                "start": start,
                "end": end,
                "line": line_no,
                "snippet": snippet,
            })
    return hits


def _check_rule_du_toan_mua_sam(text: str) -> List[Dict[str, Any]]:
    if not text:
        return []
    violations = []
    for m in RULE_BASE.finditer(text):
        start, end = m.start(), m.end()
        after = text[end:end + 50]
        if not RULE_EXPECT_AFTER.search(after):
            line_no = text.count("\n", 0, start) + 1
            left = max(0, start - 40)
            right = min(len(text), end + 60)
            snippet = text[left:right].replace("\n", " ↵ ")
            violations.append({
                "rule": "dự toán mua sắm: mua sắm",
                "message": "Sau 'dự toán mua sắm' phải có ': mua sắm'.",
                "start": start,
                "end": end,
                "line": line_no,
                "snippet": snippet,
            })
    return violations


# =========================
# Main
# =========================
def check_docx(
    file_bytes: bytes,
    phrases: List[str],
    *,
    case_sensitive: bool = False,
    whole_word: bool = False,
    scan_headers_footers: bool = True,
    spellcheck_vi: bool = False,
    spell_max_distance: int = 2,
    check_du_toan_rule: bool = False,
) -> Dict[str, Any]:
    doc = Document(io_bytes(file_bytes))

    normalized_phrases = [_safe_text(p) for p in phrases if _safe_text(p)]
    normalized_phrases = list(dict.fromkeys(normalized_phrases))

    occurrences_by_phrase: Dict[str, List[Dict[str, Any]]] = {p: [] for p in normalized_phrases}

    # We collect all text blocks with location info
    blocks: List[Dict[str, Any]] = []

    page = 1
    paragraph_index = 0
    table_index = 0

    # -------- BODY blocks (paragraphs + tables)
    for kind, block in iter_block_items(doc):
        if kind == "paragraph":
            paragraph_index += 1
            p: Paragraph = block
            txt = _safe_text(p.text)

            where = {
                "area": "body",
                "kind": "paragraph",
                "page_est": page,
                "paragraph_index": paragraph_index,
                "style_name": (p.style.name if getattr(p, "style", None) else None),
            }
            blocks.append({"text": txt, "where": where})

            matches = _scan_text_for_phrases(txt, normalized_phrases, case_sensitive=case_sensitive, whole_word=whole_word)
            for phrase, ms in matches.items():
                for item in ms:
                    occurrences_by_phrase[phrase].append({
                        "phrase": phrase,
                        "where": {**where, "line_est": item["line"]},
                        "match": {"start": item["start"], "end": item["end"]},
                        "snippet": item["snippet"],
                        "text": txt,
                    })

            if _has_page_break_in_paragraph(p):
                page += 1

        elif kind == "table":
            table_index += 1
            tbl: Table = block
            for r_i, row in enumerate(tbl.rows, start=1):
                for c_i, cell in enumerate(row.cells, start=1):
                    for cp_i, cp in enumerate(cell.paragraphs, start=1):
                        paragraph_index += 1
                        txt = _safe_text(cp.text)

                        where = {
                            "area": "body",
                            "kind": "table_cell",
                            "page_est": page,
                            "table_index": table_index,
                            "row": r_i,
                            "col": c_i,
                            "cell_paragraph_index": cp_i,
                            "global_paragraph_index": paragraph_index,
                            "style_name": (cp.style.name if getattr(cp, "style", None) else None),
                        }
                        blocks.append({"text": txt, "where": where})

                        matches = _scan_text_for_phrases(txt, normalized_phrases, case_sensitive=case_sensitive, whole_word=whole_word)
                        for phrase, ms in matches.items():
                            for item in ms:
                                occurrences_by_phrase[phrase].append({
                                    "phrase": phrase,
                                    "where": {**where, "line_est": item["line"]},
                                    "match": {"start": item["start"], "end": item["end"]},
                                    "snippet": item["snippet"],
                                    "text": txt,
                                })

                        if _has_page_break_in_paragraph(cp):
                            page += 1

    # -------- HEADERS/FOOTERS blocks
    if scan_headers_footers:
        for s_i, section in enumerate(doc.sections, start=1):
            header = section.header
            for p_i, p in enumerate(header.paragraphs, start=1):
                txt = _safe_text(p.text)
                where = {
                    "area": "header",
                    "kind": "paragraph",
                    "section": s_i,
                    "paragraph_index": p_i,
                    "page_est": None,
                    "style_name": (p.style.name if getattr(p, "style", None) else None),
                }
                blocks.append({"text": txt, "where": where})

                matches = _scan_text_for_phrases(txt, normalized_phrases, case_sensitive=case_sensitive, whole_word=whole_word)
                for phrase, ms in matches.items():
                    for item in ms:
                        occurrences_by_phrase[phrase].append({
                            "phrase": phrase,
                            "where": {**where, "line_est": item["line"]},
                            "match": {"start": item["start"], "end": item["end"]},
                            "snippet": item["snippet"],
                            "text": txt,
                        })

            footer = section.footer
            for p_i, p in enumerate(footer.paragraphs, start=1):
                txt = _safe_text(p.text)
                where = {
                    "area": "footer",
                    "kind": "paragraph",
                    "section": s_i,
                    "paragraph_index": p_i,
                    "page_est": None,
                    "style_name": (p.style.name if getattr(p, "style", None) else None),
                }
                blocks.append({"text": txt, "where": where})

                matches = _scan_text_for_phrases(txt, normalized_phrases, case_sensitive=case_sensitive, whole_word=whole_word)
                for phrase, ms in matches.items():
                    for item in ms:
                        occurrences_by_phrase[phrase].append({
                            "phrase": phrase,
                            "where": {**where, "line_est": item["line"]},
                            "match": {"start": item["start"], "end": item["end"]},
                            "snippet": item["snippet"],
                            "text": txt,
                        })

    # =========================
    # AFTER blocks built: run extra checks
    # =========================
    # A) Forbidden misspellings
    misspellings_found: List[Dict[str, Any]] = []
    # B) Rule violations (dự toán...) - optional by checkbox
    rule_violations: List[Dict[str, Any]] = []
    # C) Capitalization alerts
    capitalization_alerts: List[Dict[str, Any]] = []

    for b in blocks:
        txt = b["text"] or ""
        where = b["where"]
        if not txt.strip():
            continue

        # Forbidden words
        for item in _scan_forbidden_words(txt, FORBIDDEN_WORDS_VI):
            misspellings_found.append({
                "type": "forbidden_word",
                "where": {**where, "line_est": item["line"]},
                "wrong_word": item["wrong_word"],
                "canonical": item["canonical"],
                "match": {"start": item["start"], "end": item["end"]},
                "snippet": item["snippet"],
            })

        # Rule: dự toán mua sắm: mua sắm (only if enabled)
        if check_du_toan_rule:
            for v in _check_rule_du_toan_mua_sam(txt):
                rule_violations.append({
                    "type": "format_rule",
                    "where": {**where, "line_est": v["line"]},
                    "rule": v["rule"],
                    "message": v["message"],
                    "match": {"start": v["start"], "end": v["end"]},
                    "snippet": v["snippet"],
                })

        # Capitalization (skip heading-like lines)
        style_name = where.get("style_name")
        is_heading = _is_heading_like(txt, style_name=style_name)

        if not is_heading:
            for item in find_weird_mixed_case_words(txt):
                capitalization_alerts.append({
                    "type": "mixed_case",
                    "where": {**where, "line_est": item["line"]},
                    "word": item["word"],
                    "match": {"start": item["start"], "end": item["end"]},
                    "snippet": item["snippet"],
                    "message": item["message"],
                })

            for item in find_sentence_start_not_capitalized(txt):
                capitalization_alerts.append({
                    "type": "sentence_capitalization",
                    "where": {**where, "line_est": item["line"]},
                    "char": item["char"],
                    "match": {"start": item["start"], "end": item["end"]},
                    "snippet": item["snippet"],
                    "message": item["message"],
                })

    # Spellcheck VI fuzzy: only for missing phrases, single-word only
    typo_suspects: Dict[str, List[Dict[str, Any]]] = {}
    if spellcheck_vi:
        for phrase in normalized_phrases:
            if occurrences_by_phrase.get(phrase):
                continue
            if " " in phrase.strip():
                continue

            suspects = []
            for b in blocks:
                txt = b["text"] or ""
                if not txt:
                    continue
                near = find_vi_typos_in_block(phrase, txt, max_dist=spell_max_distance)
                for n in near:
                    suspects.append({
                        "expected": phrase,
                        "typed_in_doc": n["typed_in_doc"],
                        "distance": n["distance"],
                        "where": {**b["where"], "line_est": n["line"]},
                        "match": {"start": n["start"], "end": n["end"]},
                        "snippet": n["snippet"],
                    })

            suspects.sort(key=lambda x: (x["distance"], (x["typed_in_doc"] or "").lower()))
            if suspects:
                typo_suspects[phrase] = suspects[:50]

    # Build response hits
    found_list: List[str] = []
    missing_list: List[str] = []
    hits: List[Dict[str, Any]] = []

    for phrase in normalized_phrases:
        occs = occurrences_by_phrase.get(phrase, [])
        if occs:
            found_list.append(phrase)
        else:
            missing_list.append(phrase)

        hits.append({
            "phrase": phrase,
            "found": bool(occs),
            "count": len(occs),
            "occurrences": occs[:500],
            "typo_suspects": typo_suspects.get(phrase, []),
        })

    return {
        "stats": {
            "phrases_total": len(normalized_phrases),
            "found": len(found_list),
            "missing": len(missing_list),
            "typo_suspects_phrases": len(typo_suspects),
            "misspellings_total": len(misspellings_found),
            "rule_violations_total": len(rule_violations),
            "capitalization_total": len(capitalization_alerts),
        },
        "found": found_list,
        "missing": missing_list,
        "hits": hits,
        "alerts": {
            "misspellings_total": len(misspellings_found),
            "rule_violations_total": len(rule_violations),
            "capitalization_total": len(capitalization_alerts),
            "misspellings": misspellings_found[:500],
            "rule_violations": rule_violations[:500],
            "capitalization": capitalization_alerts[:500],
        },
        "notes": {
            "page_estimation": "page_est ước lượng theo Page Break trong DOCX (Insert → Page Break).",
            "line_estimation": "line_est ước lượng theo line break/ký tự xuống dòng có sẵn trong text.",
            "spellcheck_vi": "Spellcheck VI là fuzzy match (Damerau-Levenshtein) cho từ đơn; chỉ gợi ý 'nghi lỗi gõ sai'.",
            "capitalization": "Bỏ qua heading; cảnh báo mixed-case kiểu SƠn và chữ đầu câu sau dấu chấm không viết hoa.",
        },
    }
